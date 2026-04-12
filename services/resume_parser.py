"""Utility to extract plain text from PDF and DOCX resume files."""
import re
from pathlib import Path

# ---------------------------------------------------------------------------
# Section detection
# ---------------------------------------------------------------------------

# Maps lowercase aliases → canonical section name used as dict key
_SECTION_MAP: dict[str, str] = {
    "summary": "summary",
    "objective": "summary",
    "professional summary": "summary",
    "career summary": "summary",
    "profile": "summary",
    "about me": "summary",
    "about": "summary",
    "experience": "experience",
    "work experience": "experience",
    "professional experience": "experience",
    "employment": "experience",
    "employment history": "experience",
    "work history": "experience",
    "career history": "experience",
    "education": "education",
    "academic background": "education",
    "academic": "education",
    "qualifications": "education",
    "skills": "skills",
    "technical skills": "skills",
    "core skills": "skills",
    "key skills": "skills",
    "core competencies": "skills",
    "competencies": "skills",
    "skill set": "skills",
    "projects": "projects",
    "personal projects": "projects",
    "key projects": "projects",
    "certifications": "certifications",
    "certificates": "certifications",
    "licenses": "certifications",
    "awards": "awards",
    "achievements": "awards",
    "honors": "awards",
    "languages": "languages",
    "interests": "interests",
    "hobbies": "interests",
    "references": "references",
    "publications": "publications",
    "volunteer": "volunteer",
    "volunteering": "volunteer",
}


def _detect_header(line: str) -> tuple[bool, str]:
    """Return (is_header, canonical_name) for a single text line."""
    raw = line.strip()
    if not raw or len(raw) > 70:
        return False, ""

    # Strip common decorators: ---, ===, leading bullets
    cleaned = re.sub(r"^[-=_*•\s]+|[-=_*\s]+$", "", raw).rstrip(":").strip()
    if not cleaned:
        return False, ""

    lower = cleaned.lower()

    # Direct map lookup
    if lower in _SECTION_MAP:
        return True, _SECTION_MAP[lower]

    # All-caps short lines (e.g. "EXPERIENCE", "WORK HISTORY")
    if cleaned.isupper() and 1 <= len(cleaned.split()) <= 4:
        canonical = _SECTION_MAP.get(lower, lower)
        return True, canonical

    return False, ""


def split_sections(text: str) -> list[tuple[str, str, str]]:
    """Split resume text into named sections.

    Returns a list of ``(canonical_name, header_line, content)`` tuples:
    - ``canonical_name``: empty string for the preamble before any recognized header
    - ``header_line``: original header line including its trailing newline (empty for preamble)
    - ``content``: body text of the section
    """
    lines = text.split("\n")
    sections: list[tuple[str, str, str]] = []
    cur_name = ""
    cur_header = ""
    cur_lines: list[str] = []

    for line in lines:
        is_hdr, name = _detect_header(line)
        if is_hdr:
            sections.append((cur_name, cur_header, "\n".join(cur_lines)))
            cur_name = name
            cur_header = line + "\n"
            cur_lines = []
        else:
            cur_lines.append(line)

    sections.append((cur_name, cur_header, "\n".join(cur_lines)))
    return sections


def rebuild_from_sections(
    sections: list[tuple[str, str, str]],
    replacements: dict[str, str],
) -> str:
    """Reconstruct resume text, substituting specified section bodies.

    Args:
        sections:     Output of :func:`split_sections`.
        replacements: ``{canonical_name: new_body_text}`` — only the body is
                      replaced; the original header line is always kept.

    Returns:
        Full resume text with replaced sections.
    """
    parts: list[str] = []
    for canonical_name, header, content in sections:
        parts.append(header)
        parts.append(replacements[canonical_name] if canonical_name in replacements else content)
    return "".join(parts)


def extract_text(filepath: str) -> str:
    """Extract text from a PDF or DOCX file.

    Args:
        filepath: Absolute path to the resume file.

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: For unsupported file types.
        RuntimeError: When extraction fails.
    """
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(filepath)
    if ext == ".docx":
        return _extract_docx(filepath)
    raise ValueError(f"Unsupported file type: {ext!r}. Only PDF and DOCX are supported.")


def _extract_pdf(filepath: str) -> str:
    try:
        import PyPDF2
    except ImportError as exc:
        raise RuntimeError("PyPDF2 is not installed. Run: pip install PyPDF2") from exc

    pages = []
    with open(filepath, "rb") as fh:
        reader = PyPDF2.PdfReader(fh)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)
    return "\n".join(pages)


def _extract_docx(filepath: str) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx") from exc

    doc = docx.Document(filepath)
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def extract_hyperlinks(filepath: str) -> dict:
    """Extract embedded hyperlink URIs from a PDF or DOCX file.

    Returns a dict with keys ``linkedin`` and ``portfolio`` (both str, may be
    empty).  Email ``mailto:`` links are intentionally excluded — email is
    handled separately via the text-extraction path.

    For PDFs, URI annotations on every page are scanned.
    For DOCX, relationship targets of type hyperlink are scanned.
    A URL is classified as LinkedIn if ``linkedin.com`` appears in it;
    the first remaining non-mailto URL becomes the portfolio link.
    """
    links: dict = {"linkedin": "", "portfolio": ""}
    path = Path(filepath)
    ext  = path.suffix.lower()

    if ext == ".pdf":
        try:
            import PyPDF2
            with open(filepath, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                for page in reader.pages:
                    if "/Annots" not in page:
                        continue
                    for annot in page["/Annots"]:
                        obj    = annot.get_object()
                        action = obj.get("/A", {})
                        if action.get("/S") != "/URI":
                            continue
                        uri = action.get("/URI", "")
                        if not uri or uri.startswith("mailto:"):
                            continue
                        if "linkedin.com" in uri.lower() and not links["linkedin"]:
                            links["linkedin"] = uri
                        elif not links["portfolio"] and "linkedin.com" not in uri.lower():
                            links["portfolio"] = uri
        except Exception:
            pass  # best-effort; fall back to text-extracted values

    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(filepath)
            for rel in doc.part.rels.values():
                if "hyperlink" not in rel.reltype:
                    continue
                uri = str(rel._target)
                if not uri or uri.startswith("mailto:"):
                    continue
                if "linkedin.com" in uri.lower() and not links["linkedin"]:
                    links["linkedin"] = uri
                elif not links["portfolio"] and "linkedin.com" not in uri.lower():
                    links["portfolio"] = uri
        except Exception:
            pass  # best-effort

    return links
