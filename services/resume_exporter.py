"""Export optimized resume to PDF and DOCX formats.

Two export paths:
  1. JSON-based (preferred): export_pdf_from_json / export_docx_from_json
     Build directly from the structured resume dict — no text intermediate.
  2. Text-based (legacy fallback): export_pdf / export_docx
     Parse a plain-text string and render it.
"""
import io
import logging
import re

logger = logging.getLogger("resume_exporter")


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _split_sections(text: str) -> list[tuple[str, list[str]]]:
    """Parse plain-text resume into (heading, [lines]) pairs.

    A line is treated as a section heading when it is short (≤60 chars),
    all-caps or title-case, and followed by non-empty content.
    """
    heading_re = re.compile(
        r"^(?:[A-Z][A-Z\s&/\-]{2,}|[A-Z][a-z]+(?: [A-Z][a-z]+)*)$"
    )
    sections: list[tuple[str, list[str]]] = []
    current_heading = ""
    current_lines: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if len(line) <= 60 and heading_re.match(line):
            if current_heading or current_lines:
                sections.append((current_heading, current_lines))
            current_heading = line
            current_lines = []
        else:
            current_lines.append(line)

    if current_heading or current_lines:
        sections.append((current_heading, current_lines))

    return sections


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def export_docx(resume_text: str) -> bytes:
    """Return a DOCX file as bytes for the given plain-text resume."""
    try:
        import docx
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx") from exc

    doc = docx.Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = docx.shared.Inches(0.75)
        section.bottom_margin = docx.shared.Inches(0.75)
        section.left_margin = docx.shared.Inches(1)
        section.right_margin = docx.shared.Inches(1)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    sections = _split_sections(resume_text)

    for heading, lines in sections:
        if heading:
            h = doc.add_paragraph(heading)
            h.style = "Heading 2"
            run = h.runs[0]
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            # Underline divider via bottom border via paragraph shading is complex;
            # a simple bold heading is clean enough.
            run.bold = True

        for line in lines:
            bullet = line.startswith(("•", "-", "*", "–"))
            clean = line.lstrip("•-*– ").strip()
            if bullet:
                p = doc.add_paragraph(clean, style="List Bullet")
            else:
                p = doc.add_paragraph(clean)
            p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph("")  # spacer

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def export_pdf(resume_text: str) -> bytes:
    """Return a PDF file as bytes for the given plain-text resume."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        )
    except ImportError as exc:
        raise RuntimeError("reportlab is not installed. Run: pip install reportlab") from exc

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        "ResumeHeading",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#1E293B"),
        spaceBefore=10,
        spaceAfter=2,
        fontName="Helvetica-Bold",
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=2,
        fontName="Helvetica",
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet",
        parent=body_style,
        leftIndent=16,
        bulletIndent=6,
        spaceAfter=1,
    )

    story = []
    resume_sections = _split_sections(resume_text)

    for heading, lines in resume_sections:
        if heading:
            story.append(Paragraph(heading, heading_style))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor("#CBD5E1"), spaceAfter=4))

        for line in lines:
            is_bullet = line.startswith(("•", "-", "*", "–"))
            clean = line.lstrip("•-*– ").strip()
            # Escape XML special chars for reportlab
            clean = clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if is_bullet:
                story.append(Paragraph(f"• {clean}", bullet_style))
            else:
                story.append(Paragraph(clean, body_style))

        story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# JSON-based export helpers (single source of truth)
# ---------------------------------------------------------------------------

def _hex(s: str) -> "colors.HexColor":
    """Return a reportlab HexColor from a CSS hex string."""
    from reportlab.lib import colors as _colors
    return _colors.HexColor(s)


def _rl_esc(s: str) -> str:
    """Escape XML-special characters for reportlab Paragraph text."""
    if not isinstance(s, str):
        return ""
    return (
        s.replace("&", "&amp;")
         .replace("<", "&lt;")
         .replace(">", "&gt;")
    )


def _rl_link(url: str, color: str = "#1D4ED8") -> str:
    """Return a reportlab Paragraph markup string for a clickable hyperlink.

    Displays the URL with http(s):// stripped for readability, but the full
    URL is used as the href so it is clickable in PDF viewers.
    If ``url`` does not start with http, it is rendered as plain escaped text.
    """
    if not url:
        return ""
    if not url.lower().startswith("http"):
        return _rl_esc(url)
    display = re.sub(r"^https?://", "", url).rstrip("/")
    return f'<a href="{_rl_esc(url)}" color="{color}">{_rl_esc(display)}</a>'


def _docx_rgb(hex_str: str):
    """Return a python-docx RGBColor from a CSS hex string."""
    from docx.shared import RGBColor
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _validate_export_json(resume_json: dict, label: str) -> None:
    """Log warnings for common structural problems in the resume JSON."""
    h = resume_json.get("header", {})
    email = h.get("email", "")
    linkedin = h.get("linkedin", "")
    portfolio = h.get("portfolio", "")
    exp = resume_json.get("experience", [])

    logger.info(
        "%s: email=%r linkedin=%r portfolio=%r exp_jobs=%d certs=%d",
        label, email, linkedin, portfolio,
        len(exp),
        len(resume_json.get("certifications", [])),
    )
    if not email:
        logger.warning("%s: header.email is empty — will be missing from export", label)
    if " " in email:
        logger.warning("%s: header.email contains spaces: %r", label, email)
    if not linkedin:
        logger.warning("%s: header.linkedin is empty", label)
    for i, job in enumerate(exp):
        role = job.get("role") or job.get("title_line") or ""
        if not role:
            logger.warning("%s: experience[%d] missing role/title", label, i)
        loc = job.get("location", "") or job.get("location_period", "")
        if not loc:
            logger.debug("%s: experience[%d] location empty", label, i)
        if not job.get("start_date"):
            logger.debug("%s: experience[%d] start_date empty", label, i)
    logger.debug(
        "%s: final JSON dump:\n%s",
        label,
        __import__("json").dumps(resume_json, ensure_ascii=False, indent=2)[:2000],
    )


def export_pdf_from_json(resume_json: dict) -> bytes:
    """Build a PDF directly from the structured resume JSON dict.

    This is the canonical export path — it never passes through a text
    intermediary, so formatting is guaranteed to match the UI preview.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
        )
        from reportlab.lib.enums import TA_LEFT
    except ImportError as exc:
        raise RuntimeError("reportlab is not installed. Run: pip install reportlab") from exc

    _validate_export_json(resume_json, "export_pdf_from_json")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.70 * inch,
        bottomMargin=0.70 * inch,
    )

    st = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "RvName", parent=st["Normal"],
        fontSize=16, fontName="Helvetica-Bold",
        textColor=colors.HexColor("#1E293B"),
        spaceAfter=3,
    )
    contact_style = ParagraphStyle(
        "RvContact", parent=st["Normal"],
        fontSize=9, fontName="Helvetica",
        textColor=colors.HexColor("#475569"),
        spaceAfter=4,
    )
    section_style = ParagraphStyle(
        "RvSection", parent=st["Normal"],
        fontSize=10, fontName="Helvetica-Bold",
        textColor=colors.HexColor("#1E293B"),
        spaceBefore=10, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "RvBody", parent=st["Normal"],
        fontSize=10, fontName="Helvetica",
        textColor=colors.HexColor("#1E293B"),
        leading=14, spaceAfter=2,
    )
    job_title_style = ParagraphStyle(
        "RvJobTitle", parent=body_style,
        fontName="Helvetica-Bold", spaceAfter=1,
    )
    job_meta_style = ParagraphStyle(
        "RvJobMeta", parent=body_style,
        fontSize=9, fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#475569"), spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "RvBullet", parent=body_style,
        leftIndent=12, spaceAfter=2,
    )

    def sec_hr():
        return HRFlowable(
            width="100%", thickness=0.5,
            color=colors.HexColor("#CBD5E1"), spaceAfter=3,
        )

    story = []

    # ── Header ──────────────────────────────────────────────────────────────
    # Layout: Name (bold large) → Title (medium) → Location | Phone | Email | LinkedIn | Portfolio
    h = resume_json.get("header", {})
    if h.get("name"):
        story.append(Paragraph(_rl_esc(h["name"]), name_style))

    if h.get("title"):
        title_style_inst = ParagraphStyle(
            "RvTitle", parent=st["Normal"],
            fontSize=11, fontName="Helvetica",
            textColor=colors.HexColor("#334155"),
            spaceAfter=2,
        )
        story.append(Paragraph(_rl_esc(h["title"]), title_style_inst))

    # Contact line: location first (if present), then phone, email, links
    contact_parts = []
    for key in ("location", "phone", "email"):
        v = h.get(key, "")
        if v:
            contact_parts.append(_rl_esc(v))
    for key in ("linkedin", "portfolio"):
        v = h.get(key, "")
        if v:
            contact_parts.append(_rl_link(v))
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))

    story.append(HRFlowable(width="100%", thickness=1.5,
                             color=colors.HexColor("#1E293B"), spaceAfter=6))

    # ── Professional Summary ─────────────────────────────────────────────────
    if resume_json.get("professional_summary"):
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
        story.append(sec_hr())
        for line in resume_json["professional_summary"].split("\n"):
            if line.strip():
                story.append(Paragraph(_rl_esc(line.strip()), body_style))
        story.append(Spacer(1, 4))

    # ── Areas of Expertise ───────────────────────────────────────────────────
    if resume_json.get("areas_of_expertise"):
        story.append(Paragraph("AREAS OF EXPERTISE", section_style))
        story.append(sec_hr())
        for item in resume_json["areas_of_expertise"]:
            story.append(Paragraph(f"• {_rl_esc(item)}", bullet_style))
        story.append(Spacer(1, 4))

    # ── Professional Experience ──────────────────────────────────────────────
    if resume_json.get("experience"):
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", section_style))
        story.append(sec_hr())
        for job in resume_json["experience"]:
            # Support both new schema (role/company) and old (title_line)
            role    = job.get("role")    or job.get("title_line") or ""
            company = job.get("company") or ""
            title   = f"{role} \u2014 {company}" if company else role
            if title:
                story.append(Paragraph(_rl_esc(title), job_title_style))

            # Meta: Location | StartDate – EndDate
            loc        = job.get("location",   "") or job.get("location_period", "") or ""
            start_date = job.get("start_date", "") or ""
            end_date   = job.get("end_date",   "") or ""
            meta_parts = []
            if loc:
                meta_parts.append(loc)
            if start_date:
                dr = start_date + (f" \u2013 {end_date}" if end_date else "")
                meta_parts.append(dr)
            if meta_parts:
                story.append(Paragraph(_rl_esc(" | ".join(meta_parts)), job_meta_style))

            # Highlights / bullets
            for b in (job.get("highlights") or job.get("bullets") or []):
                story.append(Paragraph(f"• {_rl_esc(b)}", bullet_style))
            story.append(Spacer(1, 4))

    # ── Achievements & Awards ────────────────────────────────────────────────
    if resume_json.get("achievements_awards"):
        story.append(Paragraph("ACHIEVEMENTS & AWARDS", section_style))
        story.append(sec_hr())
        for item in resume_json["achievements_awards"]:
            story.append(Paragraph(f"• {_rl_esc(item)}", bullet_style))
        story.append(Spacer(1, 4))

    # ── Certifications ───────────────────────────────────────────────────────
    if resume_json.get("certifications"):
        story.append(Paragraph("CERTIFICATIONS", section_style))
        story.append(sec_hr())
        for item in resume_json["certifications"]:
            story.append(Paragraph(f"• {_rl_esc(item)}", bullet_style))
        story.append(Spacer(1, 4))

    # ── Education ────────────────────────────────────────────────────────────
    if resume_json.get("education"):
        story.append(Paragraph("EDUCATION", section_style))
        story.append(sec_hr())
        for ed in resume_json["education"]:
            if ed.get("degree"):
                story.append(Paragraph(_rl_esc(ed["degree"]), job_title_style))
            if ed.get("institution"):
                story.append(Paragraph(_rl_esc(ed["institution"]), body_style))
            if ed.get("period"):
                story.append(Paragraph(_rl_esc(ed["period"]), job_meta_style))
            story.append(Spacer(1, 3))

    doc.build(story)
    return buf.getvalue()


def export_docx_from_json(resume_json: dict) -> bytes:
    """Build a DOCX directly from the structured resume JSON dict.

    This is the canonical export path — it never passes through a text
    intermediary, so formatting is guaranteed to match the UI preview.
    """
    try:
        import docx
        from docx.shared import Pt, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    _validate_export_json(resume_json, "export_docx_from_json")

    d = docx.Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for sec in d.sections:
        sec.top_margin    = Inches(0.70)
        sec.bottom_margin = Inches(0.70)
        sec.left_margin   = Inches(0.85)
        sec.right_margin  = Inches(0.85)

    # Remove default paragraph spacing from Normal style
    d.styles["Normal"].paragraph_format.space_after  = Pt(0)
    d.styles["Normal"].paragraph_format.space_before = Pt(0)

    # ── Helpers ───────────────────────────────────────────────────────────
    def _para(text, bold=False, italic=False, size=10,
               color="#1E293B", space_after=2, indent_in=0.0) -> "docx.text.paragraph.Paragraph":
        p = d.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        if indent_in:
            p.paragraph_format.left_indent = Inches(indent_in)
        run = p.add_run(text or "")
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = _docx_rgb(color)
        return p

    def _bullet_para(text, size=10):
        return _para(f"• {text}", size=size, indent_in=0.15, space_after=1)

    def _section_title(title):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _docx_rgb("#1E293B")
        # Bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "CBD5E1")
        pBdr.append(bot)
        pPr.append(pBdr)
        return p

    def _spacer():
        d.add_paragraph()

    def _add_hyperlink(para, text: str, url: str, size: int = 9, color: str = "1D4ED8"):
        """Append a clickable OOXML hyperlink run to an existing paragraph."""
        # Register the relationship
        part = para.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Hyperlink style: blue, underline
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), color.lstrip("#"))
        rPr.append(color_el)

        u_el = OxmlElement("w:u")
        u_el.set(qn("w:val"), "single")
        rPr.append(u_el)

        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), str(size * 2))
        rPr.append(sz_el)

        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_run.append(t)
        hyperlink.append(new_run)
        para._element.append(hyperlink)

    # ── Header ────────────────────────────────────────────────────────────
    # Layout: Name (bold large) → Title → Location | Phone | Email | LinkedIn | Portfolio
    h = resume_json.get("header", {})
    if h.get("name"):
        _para(h["name"], bold=True, size=16, space_after=2)

    if h.get("title"):
        _para(h["title"], size=11, color="#334155", space_after=2)

    # Contact line: plain-text fields first, then hyperlinked fields
    plain_parts = []
    for key in ("location", "phone", "email"):
        v = h.get(key, "")
        if v:
            plain_parts.append(v)

    link_fields = []
    for key in ("linkedin", "portfolio"):
        v = h.get(key, "")
        if v:
            display = re.sub(r"^https?://", "", v).rstrip("/") if v.lower().startswith("http") else v
            link_fields.append((display, v if v.lower().startswith("http") else None))

    if plain_parts or link_fields:
        contact_p = d.add_paragraph()
        contact_p.paragraph_format.space_after  = Pt(4)
        contact_p.paragraph_format.space_before = Pt(0)
        full_plain = " | ".join(plain_parts)
        if link_fields:
            full_plain += " | " if full_plain else ""
        if full_plain:
            run = contact_p.add_run(full_plain)
            run.font.size = Pt(9)
            run.font.color.rgb = _docx_rgb("#475569")
        for idx, (display, url) in enumerate(link_fields):
            if url:
                _add_hyperlink(contact_p, display, url, size=9)
            else:
                run = contact_p.add_run(display)
                run.font.size = Pt(9)
                run.font.color.rgb = _docx_rgb("#475569")
            if idx < len(link_fields) - 1:
                sep = contact_p.add_run(" | ")
                sep.font.size = Pt(9)
                sep.font.color.rgb = _docx_rgb("#94A3B8")

    # ── Professional Summary ──────────────────────────────────────────────
    if resume_json.get("professional_summary"):
        _section_title("PROFESSIONAL SUMMARY")
        for line in resume_json["professional_summary"].split("\n"):
            if line.strip():
                _para(line.strip(), space_after=2)
        _spacer()

    # ── Areas of Expertise ────────────────────────────────────────────────
    if resume_json.get("areas_of_expertise"):
        _section_title("AREAS OF EXPERTISE")
        for item in resume_json["areas_of_expertise"]:
            _bullet_para(item)
        _spacer()

    # ── Professional Experience ───────────────────────────────────────────
    if resume_json.get("experience"):
        _section_title("PROFESSIONAL EXPERIENCE")
        for job in resume_json["experience"]:
            role    = job.get("role")    or job.get("title_line") or ""
            company = job.get("company") or ""
            title   = f"{role} \u2014 {company}" if company else role
            if title:
                _para(title, bold=True, space_after=1)

            loc        = job.get("location",   "") or job.get("location_period", "") or ""
            start_date = job.get("start_date", "") or ""
            end_date   = job.get("end_date",   "") or ""
            meta_parts = []
            if loc:
                meta_parts.append(loc)
            if start_date:
                dr = start_date + (f" \u2013 {end_date}" if end_date else "")
                meta_parts.append(dr)
            if meta_parts:
                _para(" | ".join(meta_parts), italic=True, size=9,
                      color="#475569", space_after=2)

            for b in (job.get("highlights") or job.get("bullets") or []):
                _bullet_para(b)
            _spacer()

    # ── Achievements & Awards ─────────────────────────────────────────────
    if resume_json.get("achievements_awards"):
        _section_title("ACHIEVEMENTS & AWARDS")
        for item in resume_json["achievements_awards"]:
            _bullet_para(item)
        _spacer()

    # ── Certifications ────────────────────────────────────────────────────
    if resume_json.get("certifications"):
        _section_title("CERTIFICATIONS")
        for item in resume_json["certifications"]:
            _bullet_para(item)
        _spacer()

    # ── Education ─────────────────────────────────────────────────────────
    if resume_json.get("education"):
        _section_title("EDUCATION")
        for ed in resume_json["education"]:
            if ed.get("degree"):
                _para(ed["degree"], bold=True, space_after=1)
            if ed.get("institution"):
                _para(ed["institution"], space_after=1)
            if ed.get("period"):
                _para(ed["period"], italic=True, size=9,
                      color="#475569", space_after=2)
        _spacer()

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
